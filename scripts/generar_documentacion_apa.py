"""
Genera documentación del proyecto Calculadora IMC+ en Word con formato APA 7 (estudiante).
Ejecutar: python scripts/generar_documentacion_apa.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, Inches


def set_paragraph_apa(p, bold=False, center=False, first_line_indent_pt=0):
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Pt(first_line_indent_pt)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = bold
    if not p.runs:
        pass


def add_heading(doc, text, level=1):
    """Nivel 1 centrado y negrita (APA); niveles 2-3 alineados a la izquierda."""
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
    else:
        run = p.add_run(text)
        run.bold = True
    set_paragraph_apa(p)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, text, indent_first=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    fi = 36 if indent_first else 0  # 0.5 in ≈ 36 pt
    p.paragraph_format.first_line_indent = Pt(fi)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def add_reference(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def main():
    out_dir = Path(__file__).resolve().parent.parent / "docs"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / "Documentacion_Calculadora_IMC_APA.docx"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # ----- Portada (APA 7, trabajo de estudiante) -----
    for _ in range(8):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(
        "Documentación técnica y de usuario: aplicación móvil Calculadora IMC+ "
        "(Expo y React Native)"
    )
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    for _ in range(3):
        doc.add_paragraph()

    for line in (
        "Nicolas Fajardo (ID 409483)",
        "Valentina Quinayas (ID 386959)",
        "",
        "Universidad Católica [completar sede o programa]",
        "Desarrollo de aplicaciones móviles",
        "",
        "Docente: [completar]",
        "",
        "Fecha de entrega: 15 de abril de 2026",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.space_after = Pt(0)

    doc.add_page_break()

    # ----- Resumen -----
    add_heading(doc, "Resumen", level=1)
    add_body(
        doc,
        "El presente documento consolida la documentación de la aplicación Calculadora IMC+, "
        "desarrollada con Expo (~54) y React Native, con navegación mediante Expo Router. "
        "La app calcula el índice de masa corporal (IMC) para adultos según puntos de corte "
        "habituales de la Organización Mundial de la Salud (OMS) y, para niños y adolescentes "
        "entre 5 y 19 años, utiliza el enfoque de IMC para la edad con tablas LMS de referencia "
        "OMS 2007, expresando el resultado como puntaje Z y categoría. Se describen requisitos, "
        "instalación, estructura del código, flujos de pantalla y referencias bibliográficas en "
        "formato APA 7.",
        indent_first=True,
    )
    add_body(
        doc,
        "Palabras clave: índice de masa corporal, Expo, React Native, OMS, LMS, salud infantil.",
        indent_first=False,
    )

    doc.add_page_break()

    # ----- Cuerpo -----
    add_heading(doc, "Introducción", level=1)
    add_body(
        doc,
        "Calculadora IMC+ es una aplicación móvil multiplataforma orientada a estimar el IMC "
        "a partir de peso (kg) y altura (m), diferenciando el cálculo para adultos y para "
        "niños y adolescentes. La interfaz permite elegir el perfil en la pantalla de inicio, "
        "completar los datos solicitados y visualizar categoría, recomendaciones generales de "
        "hábitos y una representación visual del resultado. El propósito de esta documentación "
        "es facilitar la comprensión del proyecto para fines académicos, despliegue y mantenimiento.",
        indent_first=True,
    )

    add_heading(doc, "Descripción general del producto", level=2)
    add_body(
        doc,
        "Nombre de la aplicación: Calculadora IMC+ (identificador de proyecto v2calculadoraimc). "
        "Versión indicada en el manifiesto: 1.0.0. La orientación predeterminada es vertical; "
        "la interfaz se adapta a tema claro u oscuro según la configuración del sistema.",
        indent_first=True,
    )

    add_heading(doc, "Requisitos del entorno de desarrollo", level=2)
    add_body(
        doc,
        "Se recomienda Node.js en versión compatible con el ecosistema Expo 54 (por ejemplo Node.js 20 LTS "
        "o superior; el repositorio público sugiere Node.js 24). Se requiere npm (incluido con Node) "
        "o un gestor alternativo (yarn, pnpm). Para prueba en dispositivo físico es opcional la app "
        "Expo Go. Para Android en emulador o compilación nativa, Android Studio y SDK configurados; "
        "para iOS, Xcode en macOS.",
        indent_first=True,
    )

    add_heading(doc, "Instalación y puesta en marcha", level=2)
    add_body(
        doc,
        "Desde la raíz del proyecto: (1) instalar dependencias con npm install; (2) iniciar el "
        "servidor de desarrollo con npx expo start o npm run start. En la terminal de Expo se "
        "puede abrir la app en web (tecla w o script npm run web), en Android (tecla a o npm run android) "
        "o en iOS (npm run ios) según el entorno disponible. Si el puerto 8081 está ocupado, "
        "puede usarse npx expo start --port 8082.",
        indent_first=True,
    )

    add_heading(doc, "Arquitectura y stack tecnológico", level=2)
    add_body(
        doc,
        "El punto de entrada principal es expo-router/entry, con rutas basadas en archivos bajo "
        "la carpeta app/. La navegación raíz es un Stack de Expo Router con pantallas index (inicio), "
        "adultos y ninos. Las dependencias principales incluyen expo (~54), react-native (~0.81), "
        "react 19, expo-router, react-native-reanimated, react-native-gesture-handler y "
        "react-native-safe-area-context, entre otras listadas en package.json.",
        indent_first=True,
    )

    add_heading(doc, "Estructura relevante del proyecto", level=2)
    add_body(
        doc,
        "app/_layout.tsx: define el Stack y títulos de cabecera. app/index.tsx: pantalla de bienvenida "
        "y selección Adultos / Niños y adolescentes. app/adultos.tsx: formulario edad, peso, altura y "
        "resultado para adultos. app/ninos.tsx: sexo biológico (curvas OMS), edad, peso, altura y "
        "resultado para 5–19 años. app/logic/calcularIMC.ts: IMC, categorías OMS para adultos, barra "
        "de posición y recomendaciones. app/logic/calcularIMCNinosOMS.ts y whoBmiLmsData.ts: método LMS, "
        "z-score e interpolación mensual. app/components/: tarjetas de resultado y elementos de interfaz. "
        "app/styles/homeStyles.ts: estilos compartidos. scripts/gen-who-lms.mjs: utilidad para datos LMS.",
        indent_first=True,
    )

    add_heading(doc, "Funcionalidades por pantalla", level=2)
    add_body(
        doc,
        "Inicio: dos accesos diferenciados con texto aclaratorio (adultos: IMC clásico y curso opcional; "
        "niños: referencia OMS 2007, 5 a 19 años). Adultos: validación de campos numéricos, cálculo de "
        "IMC, categoría (bajo peso, normal, sobrepeso, obesidad), color asociado, barra indicadora, "
        "recomendaciones contextualizadas y botón para nuevo cálculo con desplazamiento al inicio del "
        "formulario. Niños y adolescentes: selección Niño/Niña, mismos campos con edad en años "
        "(admite decimales con coma o punto), cálculo según OMS 2007 con z-score, categorías por "
        "umbrales de Z (por ejemplo delgadez severa, adecuado, sobrepeso, obesidad), texto de referencia "
        "y tarjeta de resultado específica.",
        indent_first=True,
    )

    add_heading(doc, "Consideraciones de uso y limitaciones", level=2)
    add_body(
        doc,
        "Las recomendaciones mostradas son de carácter general y educativo; no sustituyen valoración "
        "médica ni nutricional. En menores, las decisiones deben coordinarse con pediatría. Los datos "
        "de referencia OMS están embebidos en el código para funcionamiento offline; cualquier "
        "actualización oficial de tablas LMS debería revisarse y, si aplica, regenerarse con los "
        "scripts del repositorio.",
        indent_first=True,
    )

    add_heading(doc, "Tabla de comandos frecuentes", level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Comando"
    hdr[1].text = "Descripción"
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
    rows = [
        ("npm install", "Instala dependencias del proyecto."),
        ("npx expo start", "Inicia Metro y la interfaz de desarrollo de Expo."),
        ("npx expo start --tunnel", "Útil si el móvil no está en la misma red Wi-Fi."),
        ("npx expo start --clear", "Limpia caché de Metro ante errores extraños."),
        ("npm run web", "Abre la variante web si está configurada."),
    ]
    for a, b in rows:
        row = tbl.add_row().cells
        row[0].text = a
        row[1].text = b
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
    doc.add_paragraph()

    add_heading(doc, "Referencias", level=1)
    refs = [
        "Expo. (s. f.). Documentación de Expo. https://docs.expo.dev/",
        "Meta Open Source. (s. f.). React Native. https://reactnative.dev/",
        "Organización Mundial de la Salud. (s. f.). Child growth standards: BMI-for-age (5-19 years). "
        "https://www.who.int/tools/child-growth-standards/standards/body-mass-index-for-age-bmi-for-age",
        "American Psychological Association. (2020). Publication manual of the American Psychological "
        "Association (7th ed.). American Psychological Association.",
    ]
    for ref in refs:
        add_reference(doc, ref)

    doc.save(out_path)
    print(f"Generado: {out_path}")


if __name__ == "__main__":
    main()
